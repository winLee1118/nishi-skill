from __future__ import annotations

from pathlib import Path

from nihaixia_core.finalhopes_import import write_markdown
from nihaixia_core.ingest import build_index


ITEMS = [
    (
        "http://www.finalhopes.com/include/tianjishouxie.doc",
        Path("data/processed/finalhopes/tianjishouxie.docx"),
    ),
    (
        "http://www.finalhopes.com/include/ziweidoushu.doc",
        Path("data/processed/finalhopes/ziweidoushu.docx"),
    ),
]


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"表格 {table_index}\n" + "\n".join(rows))
    return "\n\n".join(parts)


def main() -> None:
    for url, path in ITEMS:
        text = extract_docx(path)
        markdown_path = write_markdown(
            url=url,
            raw_path=path,
            text=text,
            note="Converted from old .doc with Microsoft Word COM.",
            output_root=Path("knowledge/vault/_imported/finalhopes"),
            source_page="http://www.finalhopes.com/download.html",
        )
        print(markdown_path, len(text))
    print(build_index("knowledge/vault", "data/nihaixia.sqlite"))


if __name__ == "__main__":
    main()

