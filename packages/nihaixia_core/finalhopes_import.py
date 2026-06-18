from __future__ import annotations

import argparse
import html.parser
import json
import re
import shutil
import subprocess
import tempfile
import urllib.parse
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from .ingest import build_index
from .text import stable_id

DIRECT_EXTENSIONS = {".pdf", ".doc", ".docx", ".xlsx"}
DEFAULT_PAGE_URL = "http://www.finalhopes.com/download.html"


@dataclass(slots=True)
class ImportItem:
    url: str
    filename: str
    extension: str
    domain: str
    course: str
    title: str
    raw_path: Path
    markdown_path: Path
    status: str
    chars: int = 0
    note: str = ""


class LinkParser(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.links: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() != "a":
            return
        for key, value in attrs:
            if key.lower() == "href" and value:
                self.links.append(value.strip())


def fetch_text(url: str, timeout: int = 60) -> str:
    request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 nihaixia-import/0.1"})
    with urllib.request.urlopen(request, timeout=timeout) as response:
        data = response.read()
    for encoding in ("utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def discover_direct_links(page_url: str) -> list[str]:
    html = fetch_text(page_url)
    parser = LinkParser()
    parser.feed(html)
    base = urllib.parse.urlparse(page_url)
    links: list[str] = []
    seen: set[str] = set()
    for href in parser.links:
        absolute = urllib.parse.urljoin(page_url, href)
        parsed = urllib.parse.urlparse(absolute)
        extension = Path(parsed.path).suffix.lower()
        if extension not in DIRECT_EXTENSIONS:
            continue
        if parsed.netloc and parsed.netloc != base.netloc:
            continue
        normalized = urllib.parse.urlunparse(parsed._replace(fragment=""))
        if normalized not in seen:
            seen.add(normalized)
            links.append(normalized)
    return sorted(links)


def safe_filename(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    name = Path(urllib.parse.unquote(parsed.path)).name
    return re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", name) or f"{stable_id(url)}.bin"


def download(url: str, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    destination = output_dir / safe_filename(url)
    if destination.exists() and destination.stat().st_size > 0:
        return destination
    request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 nihaixia-import/0.1"})
    temp_destination = destination.with_suffix(destination.suffix + ".part")
    with urllib.request.urlopen(request, timeout=45) as response:
        with temp_destination.open("wb") as handle:
            while True:
                chunk = response.read(1024 * 256)
                if not chunk:
                    break
                handle.write(chunk)
    temp_destination.replace(destination)
    return destination


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"## Page {index}\n\n{text.strip()}")
    return "\n\n".join(pages)


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"表格 {table_index}\n" + "\n".join(rows))
    return "\n\n".join(parts)


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                rows.append(" | ".join(values).strip())
        if rows:
            parts.append(f"## Sheet: {sheet.title}\n\n" + "\n".join(rows))
    return "\n\n".join(parts)


def extract_doc_with_converter(path: Path) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return ""
    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, str(path)],
            check=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        txt_path = Path(tmp) / f"{path.stem}.txt"
        if txt_path.exists():
            return txt_path.read_text(encoding="utf-8", errors="ignore")
    return ""


def extract_text(path: Path) -> tuple[str, str]:
    extension = path.suffix.lower()
    try:
        if extension == ".pdf":
            return extract_pdf(path), ""
        if extension == ".docx":
            return extract_docx(path), ""
        if extension == ".xlsx":
            return extract_xlsx(path), ""
        if extension == ".doc":
            text = extract_doc_with_converter(path)
            if text.strip():
                return text, ""
            return "", "Old .doc parsing requires LibreOffice/soffice or antiword; downloaded but not extracted."
    except Exception as exc:  # noqa: BLE001 - import report should keep going.
        return "", f"Extraction failed: {type(exc).__name__}: {exc}"
    return "", "Unsupported extension."


def classify_by_url(url: str) -> tuple[str, str]:
    lowered = url.lower()
    if any(token in lowered for token in ("tianji", "ziweidoushu", "64ggua", "yijing", "huangji")):
        return "tianji", "天纪"
    if any(token in lowered for token in ("kanyu", "dimaidao", "diji", "yangzhai")):
        return "diji", "地纪"
    return "renji", "人纪"


def title_from_filename(filename: str) -> str:
    stem = Path(filename).stem
    known = {
        "nihandd": "汉唐中医资料",
        "zhongjing": "仲景资料",
        "nnihaixiahantangfangji": "汉唐方剂资料",
        "ziweidoushu": "紫微斗数资料",
        "tianjishouxie": "天纪手写资料",
        "duihualiangdong": "对话梁冬",
        "dimaidao": "地脉道",
        "64ggua": "六十四卦表",
    }
    return known.get(stem.lower(), stem)


def markdown_escape_yaml(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def write_markdown(
    url: str,
    raw_path: Path,
    text: str,
    note: str,
    output_root: Path,
    source_page: str,
) -> Path:
    domain, course = classify_by_url(url)
    title = title_from_filename(raw_path.name)
    source_id = f"{domain}-finalhopes-{stable_id(url, length=12)}"
    output_dir = output_root / domain
    output_dir.mkdir(parents=True, exist_ok=True)
    markdown_path = output_dir / f"{source_id}.md"
    body = text.strip()
    if not body:
        body = f"本文件已下载，但当前环境暂未解析出正文。\n\n解析备注：{note}"
    markdown = f"""---
id: {source_id}
title: "{markdown_escape_yaml(title)}"
domain: {domain}
course: "{course}"
chapter: "{markdown_escape_yaml(title)}"
topics: [{course}]
entities: ["{markdown_escape_yaml(title)}"]
source_type: {raw_path.suffix.lower().lstrip(".")}
source_url: "{markdown_escape_yaml(url)}"
source_page: "{markdown_escape_yaml(source_page)}"
raw_file: "{markdown_escape_yaml(raw_path.as_posix())}"
rights_status: unknown
---

{body}
"""
    markdown_path.write_text(markdown, encoding="utf-8")
    return markdown_path


def import_finalhopes(
    page_url: str,
    raw_dir: str | Path,
    markdown_dir: str | Path,
    db_path: str | Path | None = None,
    build: bool = False,
) -> dict[str, object]:
    raw_root = Path(raw_dir)
    markdown_root = Path(markdown_dir)
    links = discover_direct_links(page_url)
    items: list[ImportItem] = []
    for url in links:
        domain, course = classify_by_url(url)
        filename = safe_filename(url)
        extension = Path(filename).suffix.lower()
        title = title_from_filename(filename)
        try:
            raw_path = download(url, raw_root)
            text, note = extract_text(raw_path)
            markdown_path = write_markdown(url, raw_path, text, note, markdown_root, page_url)
            status = "extracted" if text.strip() else "downloaded_unextracted"
            chars = len(text)
        except Exception as exc:  # noqa: BLE001 - keep batch imports resilient.
            raw_path = raw_root / filename
            markdown_path = markdown_root / domain / f"{domain}-finalhopes-{stable_id(url, length=12)}.md"
            status = "failed"
            chars = 0
            note = f"{type(exc).__name__}: {exc}"
        items.append(
            ImportItem(
                url=url,
                filename=filename,
                extension=extension,
                domain=domain,
                course=course,
                title=title,
                raw_path=raw_path,
                markdown_path=markdown_path,
                status=status,
                chars=chars,
                note=note,
            )
        )

    build_stats = None
    if build and db_path:
        build_stats = build_index("knowledge/vault", db_path)

    return {
        "page_url": page_url,
        "discovered": len(links),
        "downloaded": len(items),
        "extracted": sum(1 for item in items if item.status == "extracted"),
        "downloaded_unextracted": sum(1 for item in items if item.status == "downloaded_unextracted"),
        "failed": sum(1 for item in items if item.status == "failed"),
        "items": [
            {
                "url": item.url,
                "filename": item.filename,
                "extension": item.extension,
                "domain": item.domain,
                "course": item.course,
                "title": item.title,
                "raw_path": item.raw_path.as_posix(),
                "markdown_path": item.markdown_path.as_posix(),
                "status": item.status,
                "chars": item.chars,
                "note": item.note,
            }
            for item in items
        ],
        "build_index": build_stats,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Download and import direct finalhopes documents.")
    parser.add_argument("--page-url", default=DEFAULT_PAGE_URL)
    parser.add_argument("--raw-dir", default="data/raw/finalhopes")
    parser.add_argument("--markdown-dir", default="knowledge/vault/_imported/finalhopes")
    parser.add_argument("--db", default="data/nihaixia.sqlite")
    parser.add_argument("--build-index", action="store_true")
    args = parser.parse_args()
    report = import_finalhopes(args.page_url, args.raw_dir, args.markdown_dir, args.db, args.build_index)
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
